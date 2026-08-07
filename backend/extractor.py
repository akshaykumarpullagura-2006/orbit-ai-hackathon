import os
import sys

def extract_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF (fitz)"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        text_content = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_content.append(page.get_text())
        return "\n".join(text_content).strip()
    except Exception as e:
        return f"[PDF Extraction Error: {str(e)}]"

def extract_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx"""
    try:
        import docx
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    full_text.append(row_text)
        return "\n".join(full_text).strip()
    except Exception as e:
        return f"[DOCX Extraction Error: {str(e)}]"

def extract_txt(file_path: str) -> str:
    """Extract plain text from TXT files"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().strip()
    except Exception as e:
        return f"[TXT Extraction Error: {str(e)}]"

def extract_ocr(file_path: str) -> str:
    """Extract text from PNG/JPG image using OCR (pytesseract/easyocr/PIL fallback)"""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        extracted = pytesseract.image_to_string(img)
        return extracted.strip() if extracted.strip() else "[OCR: No readable text detected in image]"
    except Exception:
        try:
            import easyocr
            reader = easyocr.Reader(['en'], gpu=False)
            results = reader.readtext(file_path, detail=0)
            return "\n".join(results).strip()
        except Exception as e:
            return f"[OCR Extraction Note: Image loaded ({os.path.basename(file_path)}). {str(e)}]"

def extract_document_content(file_path: str) -> dict:
    """Detect file type and perform text extraction"""
    ext = file_path.split('.')[-1].lower()
    file_name = os.path.basename(file_path)
    extracted_text = ""

    if ext == 'pdf':
        extracted_text = extract_pdf(file_path)
    elif ext == 'docx':
        extracted_text = extract_docx(file_path)
    elif ext == 'txt':
        extracted_text = extract_txt(file_path)
    elif ext in ['png', 'jpg', 'jpeg']:
        extracted_text = extract_ocr(file_path)
    else:
        extracted_text = f"[Unsupported file format .{ext}]"

    return {
        "filename": file_name,
        "file_type": ext.upper(),
        "char_count": len(extracted_text),
        "extracted_text": extracted_text
    }

if __name__ == "__main__":
    if len(sys.argv) > 1:
        target_path = sys.argv[1]
        res = extract_document_content(target_path)
        print(res["extracted_text"])
